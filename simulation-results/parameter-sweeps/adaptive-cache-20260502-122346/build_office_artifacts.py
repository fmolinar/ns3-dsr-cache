from __future__ import annotations

import csv
import re
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


RESULT_DIR = Path(__file__).resolve().parent
REPORT_MD = RESULT_DIR / "REPORT.md"
OUT_DOCX = RESULT_DIR / "Adaptive_DSR_Route_Cache_Parameter_Sweep_Report.docx"
PNG_DIR = RESULT_DIR / "graphs-png"

ACCENT = RGBColor(31, 78, 121)
SOFT = RGBColor(94, 106, 122)
GRID = "D9E2EC"


def font(size: int, bold: bool = False):
    for path in (
        "/System/Library/Fonts/Supplemental/Arial Bold.ttf" if bold else "/System/Library/Fonts/Supplemental/Arial.ttf",
        "/System/Library/Fonts/Helvetica.ttc",
    ):
        try:
            return ImageFont.truetype(path, size)
        except OSError:
            pass
    return ImageFont.load_default()


def read_csv(path: Path) -> list[dict[str, str]]:
    with path.open(newline="") as f:
        return list(csv.DictReader(f))


def draw_axes(draw, x, y, w, h, y_max, y_label=""):
    axis = "#637083"
    grid = "#E6EAF0"
    draw.line((x, y + h, x + w, y + h), fill=axis, width=2)
    draw.line((x, y, x, y + h), fill=axis, width=2)
    label_font = font(16)
    for i in range(5):
        val = y_max * i / 4
        yy = y + h - (h * i / 4)
        draw.line((x, yy, x + w, yy), fill=grid, width=1)
        draw.text((x - 58, yy - 9), f"{val:.0f}", fill="#4B5563", font=label_font)
    if y_label:
        draw.text((x, y - 30), y_label, fill="#4B5563", font=label_font)


def chart_title(draw, title, subtitle=""):
    draw.text((42, 28), title, fill="#172033", font=font(30, True))
    if subtitle:
        draw.text((42, 66), subtitle, fill="#5B6678", font=font(17))


def bar_chart(path: Path, title: str, labels: list[str], values: list[float], unit: str, colors: list[str]):
    img = Image.new("RGB", (1100, 660), "white")
    draw = ImageDraw.Draw(img)
    chart_title(draw, title, unit)
    x, y, w, h = 105, 120, 910, 410
    y_max = max(values) * 1.18
    draw_axes(draw, x, y, w, h, y_max)
    gap = 36
    bar_w = (w - gap * (len(values) + 1)) / len(values)
    for i, (label, value) in enumerate(zip(labels, values)):
        bx = x + gap + i * (bar_w + gap)
        bh = value / y_max * h
        by = y + h - bh
        draw.rounded_rectangle((bx, by, bx + bar_w, y + h), radius=6, fill=colors[i % len(colors)])
        draw.text((bx + bar_w / 2, y + h + 22), label, fill="#172033", font=font(18, True), anchor="mm")
        val_txt = f"{value:.3f}" if value < 10 else f"{value:.0f}"
        draw.text((bx + bar_w / 2, by - 18), val_txt, fill="#172033", font=font(17, True), anchor="mm")
    img.save(path)


def grouped_bar_chart(path: Path, title: str, labels: list[str], adaptive: list[float], fixed: list[float], unit: str):
    img = Image.new("RGB", (1100, 660), "white")
    draw = ImageDraw.Draw(img)
    chart_title(draw, title, unit)
    x, y, w, h = 105, 125, 910, 400
    y_max = max(adaptive + fixed) * 1.18
    draw_axes(draw, x, y, w, h, y_max)
    group_gap = 54
    inner_gap = 12
    bar_w = (w - group_gap * (len(labels) + 1)) / len(labels) / 2 - inner_gap / 2
    for i, label in enumerate(labels):
        gx = x + group_gap + i * (2 * bar_w + inner_gap + group_gap)
        for j, (series, color) in enumerate(((adaptive[i], "#2F80ED"), (fixed[i], "#E07A5F"))):
            bx = gx + j * (bar_w + inner_gap)
            bh = series / y_max * h
            by = y + h - bh
            draw.rounded_rectangle((bx, by, bx + bar_w, y + h), radius=5, fill=color)
            txt = f"{series:.2f}" if series < 10 else f"{series:.0f}"
            draw.text((bx + bar_w / 2, by - 16), txt, fill="#172033", font=font(14, True), anchor="mm")
        draw.text((gx + bar_w + inner_gap / 2, y + h + 22), label, fill="#172033", font=font(18, True), anchor="mm")
    draw.rounded_rectangle((740, 42, 1015, 89), radius=6, outline="#CBD5E1", width=1)
    draw.rectangle((760, 58, 780, 74), fill="#2F80ED")
    draw.text((790, 54), "Adaptive", fill="#172033", font=font(16))
    draw.rectangle((885, 58, 905, 74), fill="#E07A5F")
    draw.text((915, 54), "Fixed", fill="#172033", font=font(16))
    img.save(path)


def line_chart(path: Path, title: str, run_paths: dict[str, Path]):
    img = Image.new("RGB", (1200, 680), "white")
    draw = ImageDraw.Draw(img)
    chart_title(draw, title, "ReceiveRate per simulated second, kbps")
    x, y, w, h = 105, 125, 980, 405
    series = {}
    y_max = 0.0
    for label, csv_path in run_paths.items():
        rows = read_csv(csv_path)
        pts = [(int(float(r["SimulationSecond"])), float(r["ReceiveRate"])) for r in rows]
        series[label] = pts
        y_max = max(y_max, *(v for _, v in pts))
    y_max *= 1.15
    draw_axes(draw, x, y, w, h, y_max)
    for sec in [0, 50, 100, 150, 200]:
        xx = x + w * sec / 200
        draw.text((xx, y + h + 24), str(sec), fill="#4B5563", font=font(15), anchor="mm")
    colors = {"10%": "#516BEB", "50%": "#00A676", "110%": "#F59E0B", "150%": "#DC2626"}
    for label, pts in series.items():
        mapped = [(x + w * sec / 199, y + h - h * val / y_max) for sec, val in pts]
        for a, b in zip(mapped, mapped[1:]):
            draw.line((*a, *b), fill=colors[label], width=3)
    lx, ly = 805, 42
    draw.rounded_rectangle((lx, ly, lx + 280, ly + 74), radius=6, outline="#CBD5E1")
    for i, label in enumerate(series.keys()):
        yy = ly + 16 + (i // 2) * 30
        xx = lx + 18 + (i % 2) * 130
        draw.line((xx, yy, xx + 30, yy), fill=colors[label], width=4)
        draw.text((xx + 40, yy - 9), label, fill="#172033", font=font(15, True))
    img.save(path)


def generate_charts():
    PNG_DIR.mkdir(exist_ok=True)
    summary = read_csv(RESULT_DIR / "summary.csv")
    compare = read_csv(RESULT_DIR / "comparison-summary.csv")
    labels = [f"{int(r['percentage'])}%" for r in summary]
    palette = ["#516BEB", "#00A676", "#F59E0B", "#DC2626"]
    bar_chart(PNG_DIR / "total-packets.png", "Adaptive sweep: total packets", labels, [float(r["total_packets"]) for r in summary], "Packets received over 200 seconds", palette)
    bar_chart(PNG_DIR / "average-receive-rate.png", "Adaptive sweep: average receive rate", labels, [float(r["avg_all_kbps"]) for r in summary], "Average kbps across the full run", palette)
    bar_chart(PNG_DIR / "peak-receive-rate.png", "Adaptive sweep: peak receive rate", labels, [float(r["peak_receive_rate_kbps"]) for r in summary], "Peak one-second kbps", palette)
    line_chart(PNG_DIR / "receive-rate-over-time.png", "Receive rate over time", {f"{int(r['percentage'])}%": RESULT_DIR / r["label"] / "metrics.csv" for r in summary})
    grouped_bar_chart(PNG_DIR / "adaptive-vs-fixed-total-packets.png", "Adaptive vs fixed: total packets", labels, [float(r["adaptive_total_packets"]) for r in compare], [float(r["fixed_total_packets"]) for r in compare], "Packets received over 200 seconds")
    grouped_bar_chart(PNG_DIR / "adaptive-vs-fixed-average-rate.png", "Adaptive vs fixed: average receive rate", labels, [float(r["adaptive_avg_all_kbps"]) for r in compare], [float(r["fixed_avg_all_kbps"]) for r in compare], "Average kbps across the full run")


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False, size: int = 9, color: RGBColor | None = None):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(text.strip())
    run.bold = bold
    run.font.name = "Arial"
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_table_borders(table):
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = OxmlElement(f"w:{edge}")
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), "4")
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), GRID)
        borders.append(tag)
    table._tbl.tblPr.append(borders)


def style_document(doc: Document):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.08
    for name, size in (("Heading 1", 16), ("Heading 2", 14), ("Heading 3", 12)):
        style = styles[name]
        style.font.name = "Arial"
        style.font.bold = True
        style.font.size = Pt(size)
        style.font.color.rgb = ACCENT
        style.paragraph_format.space_before = Pt(12 if name == "Heading 1" else 9)
        style.paragraph_format.space_after = Pt(5)
    styles["Title"].font.name = "Arial"
    styles["Title"].font.size = Pt(22)
    styles["Title"].font.bold = True
    styles["Title"].font.color.rgb = ACCENT

    header_p = section.header.paragraphs[0]
    header_p.text = "Adaptive DSR Route Cache Experiment"
    header_p.style = styles["Header"]
    header_p.runs[0].font.size = Pt(8)
    header_p.runs[0].font.color.rgb = SOFT
    footer_p = section.footer.paragraphs[0]
    footer_p.text = "ns-3 DSR parameter sweep report"
    footer_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer_p.runs[0].font.size = Pt(8)
    footer_p.runs[0].font.color.rgb = SOFT


def add_code_block(doc: Document, lines: list[str]):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.left_indent = Inches(0.15)
    p.paragraph_format.right_indent = Inches(0.15)
    p_format = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "F4F7FB")
    p_format.append(shd)
    run = p.add_run("\n".join(lines))
    run.font.name = "Courier New"
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(28, 38, 52)


def add_image(doc: Document, alt: str, rel: str):
    png = PNG_DIR / Path(rel).name.replace(".svg", ".png")
    if png.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(png), width=Inches(6.6))
        cap = doc.add_paragraph(alt)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.runs[0].font.size = Pt(8.5)
        cap.runs[0].font.color.rgb = SOFT


def parse_table(lines: list[str]):
    rows = []
    for line in lines:
        cells = [c.strip() for c in line.strip().strip("|").split("|")]
        if all(re.fullmatch(r":?-{3,}:?", c or "") for c in cells):
            continue
        rows.append(cells)
    return rows


def add_table(doc: Document, rows: list[list[str]]):
    if not rows:
        return
    cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    table.allow_autofit = True
    set_table_borders(table)
    table_font = 8 if cols >= 7 else 9
    for r_idx, row in enumerate(rows):
        for c_idx in range(cols):
            text = row[c_idx] if c_idx < len(row) else ""
            cell = table.cell(r_idx, c_idx)
            set_cell_text(cell, text.replace("`", ""), bold=(r_idx == 0), size=table_font, color=RGBColor(255, 255, 255) if r_idx == 0 else None)
            set_cell_shading(cell, "1F4E79" if r_idx == 0 else "FFFFFF")
    doc.add_paragraph()


def add_inline_runs(paragraph, text: str):
    parts = re.split(r"(`[^`]+`|\*\*[^*]+\*\*)", text)
    for part in parts:
        if not part:
            continue
        run = paragraph.add_run(part.strip("`*"))
        run.font.name = "Arial"
        run.font.size = Pt(10.5)
        if part.startswith("`"):
            run.font.name = "Courier New"
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(26, 73, 118)
        if part.startswith("**"):
            run.bold = True


def build_docx():
    doc = Document()
    style_document(doc)
    doc.core_properties.title = "Adaptive DSR Route Cache Parameter Sweep Report"
    doc.core_properties.subject = "ns-3 Dynamic Source Routing experiment"
    doc.core_properties.author = "Codex for Francisco Molinar"

    pending_para: list[str] = []
    lines = REPORT_MD.read_text().splitlines()
    idx = 0
    in_code = False
    code_lines: list[str] = []

    def flush_para():
        nonlocal pending_para
        if pending_para:
            p = doc.add_paragraph()
            add_inline_runs(p, " ".join(pending_para))
            pending_para = []

    while idx < len(lines):
        line = lines[idx]
        if line.startswith("```"):
            if in_code:
                add_code_block(doc, code_lines)
                code_lines = []
                in_code = False
            else:
                flush_para()
                in_code = True
            idx += 1
            continue
        if in_code:
            code_lines.append(line)
            idx += 1
            continue
        if not line.strip():
            flush_para()
            idx += 1
            continue
        if line.startswith("|"):
            flush_para()
            table_lines = []
            while idx < len(lines) and lines[idx].startswith("|"):
                table_lines.append(lines[idx])
                idx += 1
            add_table(doc, parse_table(table_lines))
            continue
        img_match = re.match(r"!\[([^\]]+)\]\(([^)]+)\)", line)
        if img_match:
            flush_para()
            add_image(doc, img_match.group(1), img_match.group(2))
            idx += 1
            continue
        if line.startswith("# "):
            flush_para()
            title = doc.add_paragraph(style="Title")
            title.add_run(line[2:].strip())
            sub = doc.add_paragraph()
            sub.add_run("Generated from REPORT.md with the parameter sweep charts and comparison tables preserved.").italic = True
            sub.runs[0].font.color.rgb = SOFT
            idx += 1
            continue
        if line.startswith("## "):
            flush_para()
            doc.add_heading(line[3:].strip(), level=1)
            idx += 1
            continue
        if line.startswith("### "):
            flush_para()
            doc.add_heading(line[4:].strip(), level=2)
            idx += 1
            continue
        if line.startswith("- "):
            flush_para()
            p = doc.add_paragraph(style="List Bullet")
            add_inline_runs(p, line[2:].strip())
            idx += 1
            continue
        pending_para.append(line.strip())
        idx += 1
    flush_para()
    doc.save(OUT_DOCX)


if __name__ == "__main__":
    generate_charts()
    build_docx()
    print(OUT_DOCX)
    print(PNG_DIR)
